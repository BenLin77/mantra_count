from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import lunardate
from datetime import datetime, timedelta, date
import calendar
from app.models.ceremony import Ceremony

class CalendarGenerator:
    """行事曆生成器"""
    
    @staticmethod
    def get_lunar_date(solar_date):
        """將公曆日期轉換為藏曆日期"""
        try:
            lunar = lunardate.LunarDate.fromSolarDate(
                solar_date.year, solar_date.month, solar_date.day
            )
            return lunar
        except Exception as e:
            print(f"轉換藏曆日期時出錯: {e}")
            return None
    
    @staticmethod
    def is_lotus_day(lunar_date):
        """判斷是否為蓮師薈供日（藏曆初十）"""
        return lunar_date and lunar_date.day == 10
    
    @staticmethod
    def is_dakini_day(lunar_date):
        """判斷是否為空行母薈供日（藏曆二十五）"""
        return lunar_date and lunar_date.day == 25
    
    @staticmethod
    def get_fridays_of_month(year, month):
        """獲取指定月份的所有週五"""
        fridays = []
        cal = calendar.monthcalendar(year, month)
        for week in cal:
            # 週五是索引4
            if week[4] != 0:
                fridays.append(date(year, month, week[4]))
        return fridays
    
    @staticmethod
    def get_all_special_dates(year):
        """獲取一年中所有的特殊日期"""
        special_dates = {
            'lotus_days': [],  # 蓮師薈供日
            'dakini_days': [], # 空行母薈供日
            'fridays': []      # 週五
        }
        
        # 計算一年中每一天
        start_date = date(year, 1, 1)
        end_date = date(year, 12, 31)
        delta = timedelta(days=1)
        current_date = start_date
        
        while current_date <= end_date:
            # 獲取藏曆日期
            lunar_date = CalendarGenerator.get_lunar_date(current_date)
            
            # 檢查是否為蓮師薈供日或空行母薈供日
            if CalendarGenerator.is_lotus_day(lunar_date):
                special_dates['lotus_days'].append(current_date)
            elif CalendarGenerator.is_dakini_day(lunar_date):
                special_dates['dakini_days'].append(current_date)
            
            # 檢查是否為週五
            if current_date.weekday() == 4:  # 週五是4
                special_dates['fridays'].append(current_date)
            
            current_date += delta
        
        return special_dates
    
    @staticmethod
    def create_calendar_docx(year, filename):
        """創建行事曆 docx 檔案"""
        doc = Document()
        
        # 設置文檔為橫向
        section = doc.sections[0]
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        
        # 添加標題
        title = doc.add_heading(f'台北市噶陀十方善聯佛學會 {year - 1911} 年度共修、法會活動行程表', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加副標題
        subtitle = doc.add_paragraph(f'{year - 1911} 年度固定共修 / 年度法會行事曆')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 創建表格
        table = doc.add_table(rows=7, cols=13)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 設置表格寬度
        table.autofit = False
        
        # 設置表格列寬
        col_widths = [2.5] + [2.0] * 12  # 第一列稍寬，其餘列等寬
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)
        
        # 填充表頭
        headers = ['項目', '一月', '二月', '三月', '四月', '五月', '六月', 
                  '七月', '八月', '九月', '十月', '十一月', '十二月']
        
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # 設置表頭背景色為淺灰色
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # 設置字體
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # 填充行標題
        row_titles = [
            '大圓滿前行共修 上師文集導讀',
            '綠度母共修 上師文集導讀',
            '破瓦法共修 上師文集導讀',
            '蓮師薈供',
            '空行母薈供',
            '年度法會'
        ]
        
        for i, title in enumerate(row_titles):
            cell = table.cell(i + 1, 0)
            cell.text = title
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            # 設置字體
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # 獲取特殊日期
        special_dates = CalendarGenerator.get_all_special_dates(year)
        
        # 填充蓮師薈供日期
        lotus_row = 4
        for date_obj in special_dates['lotus_days']:
            month = date_obj.month
            cell = table.cell(lotus_row, month)
            
            # 如果單元格已有內容，則添加新行
            if cell.text:
                cell.text += f'\n{date_obj.day}日'
            else:
                cell.text = f'{date_obj.day}日'
            
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 填充空行母薈供日期
        dakini_row = 5
        for date_obj in special_dates['dakini_days']:
            month = date_obj.month
            cell = table.cell(dakini_row, month)
            
            # 如果單元格已有內容，則添加新行
            if cell.text:
                cell.text += f'\n{date_obj.day}日'
            else:
                cell.text = f'{date_obj.day}日'
            
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 使用法會輪替系統填充週五日期
        ceremony_assignments = {}
        for month in range(1, 13):
            month_assignments = Ceremony.assign_ceremonies_to_fridays(year, month)
            ceremony_assignments[month] = month_assignments
        
        # 填充法會日期
        for month in range(1, 13):
            assignments = ceremony_assignments[month]
            
            # 根據法會名稱決定行位置
            ceremony_rows = {
                '大圓滿前行共修': 1,
                '綠度母共修': 2,
                '破瓦法共修': 3
            }
            
            for assignment in assignments:
                ceremony = assignment['ceremony']
                ceremony_name = ceremony.name
                date_obj = assignment['date']
                
                if ceremony_name in ceremony_rows:
                    row = ceremony_rows[ceremony_name]
                    cell = table.cell(row, month)
                    
                    # 格式化日期顯示（包含星期）
                    weekdays = ['一', '二', '三', '四', '五', '六', '日']
                    weekday = weekdays[date_obj.weekday()]
                    date_text = f'{date_obj.day}日({weekday})'
                    
                    # 如果單元格已有內容，則添加新行
                    if cell.text:
                        cell.text += f'\n{date_text}'
                    else:
                        cell.text = date_text
                    
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        # 添加頁腳說明
        footer = doc.add_paragraph()
        footer.add_run('歡迎大家踴躍參加，學習資糧，同霑法益。')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footnote = doc.add_paragraph()
        footnote.add_run('備註：年度共修及法會暫定，若遇不可抗力因素變更，請以台北中心公告為主。')
        footnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存文檔
        doc.save(filename)
        return filename
